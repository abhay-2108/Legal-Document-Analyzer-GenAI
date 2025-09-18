import io
import re
from typing import Dict, List, Any, Optional, BinaryIO
import PyPDF2
import docx
from dataclasses import dataclass
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Represents a processed document with extracted content"""
    filename: str
    content: str
    metadata: Dict[str, Any]
    file_type: str
    page_count: int
    word_count: int
    processed_timestamp: str
    redacted_content: Optional[str] = None

class DocumentProcessor:
    """Handles document upload, text extraction, and preprocessing"""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.txt': 'text/plain'
    }
    
    def __init__(self):
        # PII patterns for redaction
        self.pii_patterns = {
            'ssn': r'\b\d{3}-\d{2}-\d{4}\b|\b\d{9}\b',
            'phone': r'\b\d{3}-\d{3}-\d{4}\b|\(\d{3}\)\s*\d{3}-\d{4}',
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'credit_card': r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}\b',
            'address': r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Place|Pl)\b',
            'zip_code': r'\b\d{5}(?:-\d{4})?\b',
        }
        
        # Common legal document structure patterns
        self.structure_patterns = {
            'title': r'^[A-Z\s]{3,}$',
            'section_header': r'^\s*(?:SECTION|Section)\s+\d+',
            'article_header': r'^\s*(?:ARTICLE|Article)\s+\d+',
            'clause_number': r'^\s*\d+\.\d*\s+',
            'subsection': r'^\s*\([a-zA-Z0-9]+\)\s*'
        }

    def extract_text_from_file(self, file_content: bytes, filename: str) -> ProcessedDocument:
        """
        Extract text from uploaded file based on file type
        """
        file_extension = Path(filename).suffix.lower()
        
        if file_extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text, metadata = self._extract_from_pdf(file_content)
        elif file_extension in ['.docx']:
            text, metadata = self._extract_from_docx(file_content)
        elif file_extension == '.txt':
            text, metadata = self._extract_from_txt(file_content)
        else:
            raise ValueError(f"Text extraction not implemented for {file_extension}")
        
        # Clean and normalize text
        cleaned_text = self._clean_text(text)
        
        # Calculate statistics
        word_count = len(cleaned_text.split())
        
        return ProcessedDocument(
            filename=filename,
            content=cleaned_text,
            metadata=metadata,
            file_type=file_extension,
            page_count=metadata.get('page_count', 1),
            word_count=word_count,
            processed_timestamp=self._get_timestamp()
        )

    def _extract_from_pdf(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            # Extract metadata
            metadata = {
                'page_count': page_count,
                'pdf_info': {
                    'title': getattr(pdf_reader.metadata, 'title', None) if pdf_reader.metadata else None,
                    'author': getattr(pdf_reader.metadata, 'author', None) if pdf_reader.metadata else None,
                    'creator': getattr(pdf_reader.metadata, 'creator', None) if pdf_reader.metadata else None,
                }
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    def _extract_from_docx(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract metadata
            metadata = {
                'page_count': 1,  # Approximate for DOCX
                'paragraph_count': paragraph_count,
                'docx_info': {
                    'title': doc.core_properties.title,
                    'author': doc.core_properties.author,
                    'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None,
                }
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    def _extract_from_txt(self, file_content: bytes) -> tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Unable to decode text file with supported encodings")
            
            metadata = {
                'page_count': 1,
                'encoding': encoding,
                'line_count': text.count('\n')
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from TXT: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR/extraction issues
        text = text.replace('\x0c', '\n')  # Form feed to newline
        text = text.replace('\xa0', ' ')   # Non-breaking space to regular space
        
        # Remove page markers
        text = re.sub(r'\n--- Page \d+ ---\n', '\n', text)
        
        # Normalize line breaks
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        # Remove multiple consecutive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text

    def redact_pii(self, text: str, redaction_level: str = 'partial') -> tuple[str, List[Dict[str, Any]]]:
        """
        Redact personally identifiable information from text
        
        Args:
            text: Input text to redact
            redaction_level: 'none', 'partial', 'full'
        
        Returns:
            Tuple of (redacted_text, list_of_redactions)
        """
        if redaction_level == 'none':
            return text, []
        
        redacted_text = text
        redactions = []
        
        for pii_type, pattern in self.pii_patterns.items():
            matches = re.finditer(pattern, text, re.IGNORECASE)
            
            for match in matches:
                original_value = match.group()
                
                if redaction_level == 'partial':
                    # Partially redact (show first/last characters)
                    if pii_type == 'email':
                        # For emails: show first char and domain
                        email_parts = original_value.split('@')
                        if len(email_parts) == 2:
                            redacted_value = f"{email_parts[0][0]}***@{email_parts[1]}"
                        else:
                            redacted_value = "[EMAIL-REDACTED]"
                    elif pii_type in ['ssn', 'credit_card']:
                        # Show last 4 digits
                        digits_only = re.sub(r'\D', '', original_value)
                        redacted_value = f"***-**-{digits_only[-4:]}" if len(digits_only) >= 4 else "[REDACTED]"
                    elif pii_type == 'phone':
                        # Show area code
                        digits_only = re.sub(r'\D', '', original_value)
                        if len(digits_only) >= 10:
                            redacted_value = f"({digits_only[:3]}) ***-****"
                        else:
                            redacted_value = "[PHONE-REDACTED]"
                    else:
                        redacted_value = f"[{pii_type.upper()}-REDACTED]"
                else:  # full redaction
                    redacted_value = f"[{pii_type.upper()}-REDACTED]"
                
                redacted_text = redacted_text.replace(original_value, redacted_value, 1)
                
                redactions.append({
                    'type': pii_type,
                    'original_value': original_value,
                    'redacted_value': redacted_value,
                    'position': match.start(),
                    'length': len(original_value)
                })
        
        return redacted_text, redactions

    def identify_document_structure(self, text: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        Identify document structure elements like sections, headings, etc.
        """
        structure = {
            'titles': [],
            'sections': [],
            'articles': [],
            'clauses': [],
            'subsections': []
        }
        
        lines = text.split('\n')
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check for different structure patterns
            for element_type, pattern in self.structure_patterns.items():
                if re.match(pattern, line):
                    if element_type == 'title':
                        structure['titles'].append({
                            'text': line,
                            'line_number': line_num + 1,
                            'level': 1
                        })
                    elif element_type in ['section_header', 'article_header']:
                        key = 'sections' if 'section' in element_type else 'articles'
                        structure[key].append({
                            'text': line,
                            'line_number': line_num + 1,
                            'level': 2
                        })
                    elif element_type in ['clause_number', 'subsection']:
                        key = 'clauses' if 'clause' in element_type else 'subsections'
                        structure[key].append({
                            'text': line,
                            'line_number': line_num + 1,
                            'level': 3 if element_type == 'clause_number' else 4
                        })
        
        return structure

    def extract_key_sections(self, text: str) -> Dict[str, str]:
        """
        Extract key sections commonly found in legal documents
        """
        sections = {}
        text_lower = text.lower()
        
        # Common legal section patterns
        section_patterns = {
            'definitions': [
                r'definitions?\s*:?\s*\n',
                r'terms?\s+and\s+definitions?\s*:?\s*\n',
                r'interpretation\s*:?\s*\n'
            ],
            'parties': [
                r'parties?\s*:?\s*\n',
                r'contracting\s+parties?\s*:?\s*\n'
            ],
            'termination': [
                r'termination\s*:?\s*\n',
                r'end\s+of\s+agreement\s*:?\s*\n',
                r'expiration\s*:?\s*\n'
            ],
            'liability': [
                r'liability\s*:?\s*\n',
                r'limitation\s+of\s+liability\s*:?\s*\n',
                r'damages\s*:?\s*\n'
            ],
            'governing_law': [
                r'governing\s+law\s*:?\s*\n',
                r'applicable\s+law\s*:?\s*\n',
                r'jurisdiction\s*:?\s*\n'
            ],
            'dispute_resolution': [
                r'dispute\s+resolution\s*:?\s*\n',
                r'arbitration\s*:?\s*\n',
                r'mediation\s*:?\s*\n'
            ]
        }
        
        for section_name, patterns in section_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, text_lower)
                if match:
                    # Extract text from this section (approximately)
                    start_pos = match.start()
                    # Find the next section or end of document
                    next_section_pos = len(text)
                    
                    # Look for the next major section
                    for other_patterns in section_patterns.values():
                        for other_pattern in other_patterns:
                            other_match = re.search(other_pattern, text_lower[start_pos + 100:])
                            if other_match:
                                next_section_pos = min(next_section_pos, 
                                                     start_pos + 100 + other_match.start())
                    
                    section_text = text[start_pos:next_section_pos].strip()
                    sections[section_name] = section_text[:2000]  # Limit section size
                    break
        
        return sections

    def validate_document_format(self, text: str) -> Dict[str, Any]:
        """
        Validate if the document appears to be a legal document
        """
        validation = {
            'is_legal_document': False,
            'confidence_score': 0.0,
            'indicators': [],
            'missing_elements': [],
            'document_type_hints': []
        }
        
        text_lower = text.lower()
        
        # Legal document indicators
        legal_indicators = [
            ('contract_terms', ['agreement', 'contract', 'party', 'parties']),
            ('legal_language', ['whereas', 'hereby', 'hereafter', 'aforementioned']),
            ('obligations', ['shall', 'must', 'required', 'obligation']),
            ('legal_concepts', ['liability', 'damages', 'breach', 'terminate']),
            ('formal_structure', ['section', 'article', 'clause', 'subsection']),
            ('execution', ['signature', 'executed', 'effective date'])
        ]
        
        score = 0
        for category, terms in legal_indicators:
            found_terms = [term for term in terms if term in text_lower]
            if found_terms:
                validation['indicators'].append({
                    'category': category,
                    'found_terms': found_terms,
                    'weight': len(found_terms) / len(terms)
                })
                score += len(found_terms) / len(terms)
        
        validation['confidence_score'] = min(score / len(legal_indicators), 1.0)
        validation['is_legal_document'] = validation['confidence_score'] > 0.3
        
        # Suggest document type
        if 'lease' in text_lower or 'rent' in text_lower:
            validation['document_type_hints'].append('lease_agreement')
        elif 'employment' in text_lower or 'employee' in text_lower:
            validation['document_type_hints'].append('employment_contract')
        elif 'service' in text_lower and 'agreement' in text_lower:
            validation['document_type_hints'].append('service_agreement')
        elif 'non-disclosure' in text_lower or 'confidential' in text_lower:
            validation['document_type_hints'].append('nda')
        else:
            validation['document_type_hints'].append('general_contract')
        
        return validation

    def _get_timestamp(self) -> str:
        """Get current timestamp as ISO string"""
        from datetime import datetime
        return datetime.utcnow().isoformat()

    def get_processing_statistics(self, processed_doc: ProcessedDocument) -> Dict[str, Any]:
        """
        Generate processing statistics for the document
        """
        text = processed_doc.content
        
        return {
            'character_count': len(text),
            'word_count': len(text.split()),
            'sentence_count': len(re.findall(r'[.!?]+', text)),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
            'average_sentence_length': len(text.split()) / max(len(re.findall(r'[.!?]+', text)), 1),
            'readability_indicators': {
                'avg_word_length': sum(len(word) for word in text.split()) / max(len(text.split()), 1),
                'complex_word_ratio': len([w for w in text.split() if len(w) > 6]) / max(len(text.split()), 1)
            }
        }
